"""Resume parsing (PDF/DOCX → text) and export (text → PDF/DOCX)."""
from __future__ import annotations
import io


def parse_resume(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _parse_pdf(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    # plaintext fallback
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _parse_pdf(data: bytes) -> str:
    from pdfminer.high_level import extract_text
    return extract_text(io.BytesIO(data)) or ""


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def to_docx(text: str, title: str = "Document") -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        # treat ALL-CAPS short lines or '## ' as headings
        if line.startswith("## ") or (line.isupper() and len(line) < 40):
            doc.add_heading(line.replace("## ", ""), level=1)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(text: str, title: str = "Document") -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER, leftMargin=0.8 * inch, rightMargin=0.8 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                          fontSize=10.5, leading=14)
    head = ParagraphStyle("head", parent=styles["Heading2"], fontName="Helvetica-Bold")
    flow = []
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            flow.append(Spacer(1, 6)); continue
        if line.startswith("## ") or (line.isupper() and len(line) < 40):
            flow.append(Paragraph(line.replace("## ", ""), head))
        else:
            flow.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), body))
    doc.build(flow)
    return buf.getvalue()
